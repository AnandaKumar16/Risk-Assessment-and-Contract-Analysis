"""
Document Parser Module
Handles extraction of text from PDF, DOCX, and TXT files
"""

import re
from typing import Dict, Any, Optional
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parses various document formats and extracts text for analysis.
    Supports: PDF, DOCX, TXT
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def parse_document(self, file_path: str = None, file_bytes: bytes = None, 
                      filename: str = None) -> Dict[str, Any]:
        """
        Parse a document and extract text content.
        
        Args:
            file_path: Path to the document file (optional)
            file_bytes: Raw bytes of the document (optional, for uploaded files)
            filename: Name of the file (required if using file_bytes)
        
        Returns:
            Dictionary containing:
                - raw_text: Extracted text
                - metadata: Document metadata
                - error: Error message if parsing failed
        """
        try:
            if file_path:
                return self._parse_from_path(file_path)
            elif file_bytes and filename:
                return self._parse_from_bytes(file_bytes, filename)
            else:
                return {
                    'raw_text': '',
                    'metadata': {},
                    'error': 'Either file_path or (file_bytes + filename) must be provided'
                }
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            return {
                'raw_text': '',
                'metadata': {},
                'error': str(e)
            }
    
    def _parse_from_path(self, file_path: str) -> Dict[str, Any]:
        """Parse document from file path."""
        path = Path(file_path)
        
        if not path.exists():
            return {
                'raw_text': '',
                'metadata': {},
                'error': f'File not found: {file_path}'
            }
        
        file_ext = path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            return {
                'raw_text': '',
                'metadata': {},
                'error': f'Unsupported file format: {file_ext}'
            }
        
        with open(file_path, 'rb') as f:
            file_bytes = f.read()
        
        return self._parse_from_bytes(file_bytes, path.name)
    
    def _parse_from_bytes(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse document from bytes."""
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_bytes, filename)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_bytes, filename)
        elif file_ext == '.txt':
            return self._parse_txt(file_bytes, filename)
        else:
            return {
                'raw_text': '',
                'metadata': {},
                'error': f'Unsupported file format: {file_ext}'
            }
    
    def _parse_pdf(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF."""
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            raw_text = '\n\n'.join(text_content)
            
            # Extract metadata
            metadata = {
                'filename': filename,
                'format': 'PDF',
                'pages': len(pdf_reader.pages),
                'size_bytes': len(file_bytes)
            }
            
            # Try to get PDF metadata
            if pdf_reader.metadata:
                metadata.update({
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'creation_date': pdf_reader.metadata.get('/CreationDate', '')
                })
            
            return {
                'raw_text': raw_text,
                'metadata': metadata,
                'error': None
            }
            
        except ImportError:
            return {
                'raw_text': '',
                'metadata': {},
                'error': 'PyPDF2 library not installed. Please install: pip install PyPDF2'
            }
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            return {
                'raw_text': '',
                'metadata': {},
                'error': f'PDF parsing error: {str(e)}'
            }
    
    def _parse_docx(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX."""
        try:
            from docx import Document
            from io import BytesIO
            
            docx_file = BytesIO(file_bytes)
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            raw_text = '\n\n'.join(text_content)
            
            # Metadata
            metadata = {
                'filename': filename,
                'format': 'DOCX',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'size_bytes': len(file_bytes)
            }
            
            # Try to get document properties
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'created': str(core_props.created) if core_props.created else ''
                })
            except:
                pass
            
            return {
                'raw_text': raw_text,
                'metadata': metadata,
                'error': None
            }
            
        except ImportError:
            return {
                'raw_text': '',
                'metadata': {},
                'error': 'python-docx library not installed. Please install: pip install python-docx'
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return {
                'raw_text': '',
                'metadata': {},
                'error': f'DOCX parsing error: {str(e)}'
            }
    
    def _parse_txt(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from TXT."""
        try:
            # Try UTF-8 first
            try:
                raw_text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                raw_text = file_bytes.decode('latin-1')
            
            metadata = {
                'filename': filename,
                'format': 'TXT',
                'size_bytes': len(file_bytes),
                'lines': len(raw_text.split('\n'))
            }
            
            return {
                'raw_text': raw_text,
                'metadata': metadata,
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT: {str(e)}")
            return {
                'raw_text': '',
                'metadata': {},
                'error': f'TXT parsing error: {str(e)}'
            }
    
    def preprocess_text(self, text: str) -> str:
        """
        Clean and normalize text for analysis.
        
        Args:
            text: Raw text to preprocess
        
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with analysis
        # But preserve important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\-\(\)\[\]\'\"\n\r]', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        # Remove multiple consecutive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def detect_language(self, text: str) -> str:
        """
        Simple heuristic language detection (English vs Hindi).
        
        Args:
            text: Text to analyze
        
        Returns:
            'hindi', 'english', or 'mixed'
        """
        if not text:
            return 'unknown'
        
        # Count Devanagari characters (Hindi)
        hindi_chars = len(re.findall(r'[\u0900-\u097F]', text))
        
        # Count Latin characters (English)
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        
        total_chars = hindi_chars + english_chars
        
        if total_chars == 0:
            return 'unknown'
        
        hindi_ratio = hindi_chars / total_chars
        
        if hindi_ratio > 0.7:
            return 'hindi'
        elif hindi_ratio < 0.3:
            return 'english'
        else:
            return 'mixed'
